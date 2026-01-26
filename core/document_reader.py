"""
Document Reader Utility
Functions for reading various document types
"""

import json
import os
from pathlib import Path
from typing import List, Tuple, Optional
import pandas as pd


# Supported file extensions and their types
SUPPORTED_EXTENSIONS = {
    '.txt': 'text',
    '.md': 'text',
    '.json': 'json',
    '.csv': 'csv',
    '.html': 'text',
    '.htm': 'text',
    '.xml': 'text',
    '.pdf': 'pdf',
    '.docx': 'docx',
    '.doc': 'doc',
}


def read_text_file(file_path: str) -> str:
    """
    Read a text file with multiple encoding fallbacks.

    Args:
        file_path: Path to the text file

    Returns:
        File contents as string
    """
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode file with any supported encoding: {file_path}")


def read_pdf(file_path: str) -> str:
    """
    Read text from a PDF file.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text from all pages
    """
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError("PDF support requires pypdf. Install with: pip install pypdf")


def read_docx(file_path: str) -> str:
    """
    Read text from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text from all paragraphs
    """
    try:
        from docx import Document
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except ImportError:
        raise ImportError("DOCX support requires python-docx. Install with: pip install python-docx")


def read_document(file_path: str) -> Tuple[str, str]:
    """
    Read a document and return its content and type.

    Args:
        file_path: Path to the document

    Returns:
        Tuple of (content, file_type)
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    file_type = SUPPORTED_EXTENSIONS[ext]

    if file_type == 'text':
        content = read_text_file(file_path)
    elif file_type == 'json':
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        content = json.dumps(data, indent=2)
    elif file_type == 'csv':
        df = pd.read_csv(file_path)
        content = df.to_string()
    elif file_type == 'pdf':
        content = read_pdf(file_path)
    elif file_type in ['docx', 'doc']:
        content = read_docx(file_path)
    else:
        content = read_text_file(file_path)

    return content, file_type


def read_uploaded_file(uploaded_file) -> str:
    """
    Read content from a Streamlit uploaded file.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        File contents as string
    """
    file_name = uploaded_file.name
    ext = Path(file_name).suffix.lower()

    if ext in ['.txt', '.md', '.html', '.htm', '.xml']:
        return uploaded_file.read().decode('utf-8', errors='ignore')
    elif ext == '.json':
        data = json.load(uploaded_file)
        return json.dumps(data, indent=2)
    elif ext == '.csv':
        df = pd.read_csv(uploaded_file)
        return df.to_string()
    elif ext == '.pdf':
        try:
            import pypdf
            reader = pypdf.PdfReader(uploaded_file)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("PDF support requires pypdf. Install with: pip install pypdf")
    elif ext in ['.docx', '.doc']:
        try:
            from docx import Document
            doc = Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ImportError("DOCX support requires python-docx. Install with: pip install python-docx")
    else:
        return uploaded_file.read().decode('utf-8', errors='ignore')


def get_files_from_folder(folder_path: str, extensions: Optional[List[str]] = None) -> List[str]:
    """
    Get all matching files from a folder.

    Args:
        folder_path: Path to the folder to scan
        extensions: List of extensions to include (e.g., ['.txt', '.pdf'])
                   If None, includes all supported extensions

    Returns:
        Sorted list of file paths
    """
    if extensions is None:
        extensions = list(SUPPORTED_EXTENSIONS.keys())

    folder = Path(folder_path)
    if not folder.exists():
        raise ValueError(f"Folder does not exist: {folder_path}")
    if not folder.is_dir():
        raise ValueError(f"Path is not a directory: {folder_path}")

    files = []
    for file_path in folder.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            files.append(str(file_path))

    return sorted(files)


def get_supported_extensions() -> List[str]:
    """Get list of supported file extensions."""
    return list(SUPPORTED_EXTENSIONS.keys())


def get_extension_display_groups() -> dict:
    """Get extension groups for UI display."""
    return {
        "Text": ['.txt', '.md'],
        "PDF": ['.pdf'],
        "Word": ['.docx', '.doc'],
        "Data": ['.json', '.csv'],
        "Web": ['.html', '.htm', '.xml']
    }
